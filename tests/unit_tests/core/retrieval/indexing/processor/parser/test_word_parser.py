# Copyright (c) Huawei Technologies Co., Ltd. 2025-2025. All rights reserved.
"""
Word file parser test cases
"""

from functools import partial
import os
import tempfile
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from docx import Document as DocxDocument
from PIL import Image

from openjiuwen.core.retrieval import WordParser


class TestWordParser:
    """Word file parser tests"""

    @staticmethod
    def test_init():
        """Test initialization"""
        parser = WordParser()
        assert parser is not None

    @pytest.mark.asyncio
    async def test_parse_docx_success(self):
        """Test parsing DOCX file successfully"""
        parser = WordParser()

        # Use a real Document so iter_inner_content() returns real Paragraph objects
        doc = DocxDocument()
        doc.add_paragraph("Paragraph 1")
        doc.add_paragraph("Paragraph 2")
        blocks = list(doc.iter_inner_content())

        with patch(
            "openjiuwen.core.retrieval.indexing.processor.parser.word_parser.asyncio.to_thread"
        ) as mock_to_thread:
            mock_to_thread.side_effect = [doc, blocks]

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                temp_path = f.name

            try:
                documents = await parser.parse(temp_path, doc_id="doc_1")
                assert len(documents) == 1
                assert documents[0].id_ == "doc_1"
                assert "Paragraph 1" in documents[0].text
                assert "Paragraph 2" in documents[0].text
            finally:
                if os.path.exists(temp_path):
                    os.unlink(temp_path)

    @pytest.mark.asyncio
    async def test_parse_docx_e2e(self, save_result: bool = False):
        """End-to-end: create a real DOCX with python-docx, parse it with WordParser (no mocks)."""
        parser = WordParser()

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name

        try:
            # Create a real DOCX file with paragraphs and a table
            doc = DocxDocument()
            doc.add_heading("Hello")
            doc.add_paragraph("First paragraph.")
            doc.add_paragraph("Second paragraph with more text.")
            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "A1"
            table.rows[0].cells[1].text = "B1"
            table.rows[1].cells[0].text = "A2"
            table.rows[1].cells[1].text = "B2"
            doc.add_paragraph("Paragraph after table.")
            doc.save(temp_path)

            documents = await parser.parse(temp_path, doc_id="e2e_doc")

            assert len(documents) == 1
            assert documents[0].id_ == "e2e_doc"
            text = documents[0].text
            assert "## Hello" in text  # Heading 1 rendered as markdown
            assert "First paragraph." in text
            assert "Second paragraph with more text." in text
            assert "A1" in text and "B1" in text and "A2" in text and "B2" in text
            assert "Paragraph after table." in text
            if save_result:
                doc.save("./testdoc.docx")
                with open("./testdoc.md", "w", encoding="utf-8") as f:
                    f.write(text)
        finally:
            if os.path.exists(temp_path):
                os.unlink(temp_path)

    @pytest.mark.asyncio
    async def test_parse_docx_headings(self):
        """Test that Title and Heading 1–9 are output as markdown headings."""
        parser = WordParser()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name
        try:
            doc = DocxDocument()
            doc.add_paragraph("Document Title Here", style="Title")
            doc.add_paragraph("Section", style="Heading 1")
            doc.add_paragraph("Subsection", style="Heading 2")
            doc.add_paragraph("Normal paragraph.")
            doc.save(temp_path)
            documents = await parser.parse(temp_path, doc_id="headings_doc")
            assert len(documents) == 1
            text = documents[0].text
            assert "# Document Title Here" in text
            assert "## Section" in text
            assert "### Subsection" in text
            assert "Normal paragraph." in text
        finally:
            if os.path.exists(temp_path):
                os.unlink(temp_path)

    @pytest.mark.asyncio
    async def test_parse_docx_e2e_with_image(self, save_result: bool = False):
        """End-to-end: DOCX with an embedded image; parser extracts it and caption is included."""
        parser = WordParser()
        with tempfile.TemporaryDirectory() as tmpdir:
            img_path = os.path.join(tmpdir, "test_image.png")
            Image.new("RGB", (10, 10), color="red").save(img_path)
            img_path2 = os.path.join(tmpdir, "test_image2.png")
            Image.new("RGBA", (10, 10), color="blue").save(img_path2)
            docx_path = os.path.join(tmpdir, "with_image.docx")
            doc = DocxDocument()
            doc.add_paragraph("Before image.")
            doc.add_picture(img_path)
            doc.add_paragraph("Between image.")
            doc.add_picture(img_path2)
            doc.add_paragraph("After image.")
            doc.save(docx_path)
            fake_caption = "A red square test image."
            with patch(
                "openjiuwen.core.retrieval.indexing.processor.parser.word_parser.ImageCaptioner"
            ) as mock_captioner_cls:
                mock_captioner = MagicMock()
                mock_captioner.caption_images = AsyncMock(return_value=[fake_caption])
                mock_captioner_cls.return_value = mock_captioner
                documents = await parser.parse(docx_path, doc_id="e2e_image_doc")
            assert len(documents) == 1
            assert documents[0].id_ == "e2e_image_doc"
            text = documents[0].text
            assert f"Before image.\n{fake_caption}\nBetween image.\n{fake_caption}\nAfter image." in text
            mock_captioner.caption_images.assert_called()
            call_args = mock_captioner.caption_images.call_args[0][0]
            assert len(call_args) == 1
            assert call_args[0].endswith(".png")
            if save_result:
                doc.save("./test_doc_with_img.docx")
                with open("./test_doc_with_img.md", "w", encoding="utf-8") as f:
                    f.write(text)

    @pytest.mark.asyncio
    async def test_parse_docx_empty_document(self):
        """Test parsing empty document"""
        parser = WordParser()

        mock_doc = MagicMock()
        mock_doc.iter_inner_content = partial(iter, [])

        with patch(
            "openjiuwen.core.retrieval.indexing.processor.parser.word_parser.asyncio.to_thread"
        ) as mock_to_thread:
            mock_to_thread.side_effect = [
                mock_doc,
                [],
            ]

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                temp_path = f.name

            try:
                documents = await parser.parse(temp_path, doc_id="doc_1")
                # Empty document should return empty list
                assert len(documents) == 0
            finally:
                if os.path.exists(temp_path):
                    os.unlink(temp_path)

    @pytest.mark.asyncio
    async def test_parse_docx_file_not_found(self):
        """Test parsing non-existent file"""
        parser = WordParser()
        documents = await parser.parse("nonexistent.docx", doc_id="doc_1")
        # Should return empty list (exception is caught)
        assert len(documents) == 0

    @pytest.mark.asyncio
    async def test_parse_docx_with_exception(self):
        """Test exception during parsing"""
        parser = WordParser()

        with patch(
            "openjiuwen.core.retrieval.indexing.processor.parser.word_parser.asyncio.to_thread"
        ) as mock_to_thread:
            mock_to_thread.side_effect = Exception("DOCX parsing error")

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                temp_path = f.name

            try:
                documents = await parser.parse(temp_path, doc_id="doc_1")
                # Should return empty list (exception is caught)
                assert len(documents) == 0
            finally:
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
