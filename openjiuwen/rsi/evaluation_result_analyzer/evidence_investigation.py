# coding: utf-8
# Copyright (c) Huawei Technologies Co., Ltd. 2026. All rights reserved.
"""Bounded, read-only evidence acquisition for causal RSI diagnosis."""

from __future__ import annotations

import ast
import hashlib
import json
import math
import os
import re
import zipfile
from collections.abc import Mapping, Sequence
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

from openjiuwen.rsi.evaluation_result_analyzer.case_reader import CaseAnalysisInput

_ALLOWED_OPERATIONS = {
    "check_relation",
    "compare_numeric_change",
    "compare_runs",
    "inspect_artifact",
    "inspect_evaluation",
    "read_artifact_window",
    "read_repository_file",
    "read_event",
    "search_repository",
    "search_trace",
}
_MAX_HYPOTHESES = 6
_MAX_REQUESTS = 12
_MAX_SEARCH_RESULTS = 5
_MAX_EVENT_CHARS = 12_000
_MAX_ARTIFACT_FILE_CHARS = 200_000
_MAX_ARTIFACT_FILES = 100
_MAX_ARTIFACT_WINDOW_CHARS = 12_000
_MAX_AUTOMATIC_ARTIFACT_WINDOWS = 12
_MAX_AUTOMATIC_ARTIFACT_SOURCES = 4
_MAX_AUTOMATIC_ARTIFACT_CHARS = 96_000
_MAX_STRUCTURED_ARTIFACTS_PER_REQUEST = 16
_MAX_REPOSITORY_FILES = 2_000
_MAX_STRUCTURED_CELLS = 20_000
_MAX_STRUCTURED_PAGES = 100
_TEXT_ARTIFACT_SUFFIXES = {
    ".csv",
    ".diff",
    ".json",
    ".log",
    ".md",
    ".patch",
    ".txt",
    ".xml",
    ".yaml",
    ".yml",
}
_STRUCTURED_ARTIFACT_SUFFIXES = {".docx", ".pdf", ".pptx", ".xlsx"}
_STRUCTURED_QUERY_SUFFIX_HINTS = {
    ".docx": ("docx", "word", "paragraph", "document", "文档", "段落"),
    ".pdf": ("pdf", "page", "document", "文档", "页面"),
    ".pptx": ("pptx", "powerpoint", "slide", "deck", "presentation", "幻灯片", "演示"),
    ".xlsx": (
        "xlsx",
        "excel",
        "workbook",
        "worksheet",
        "spreadsheet",
        "cell",
        "formula",
        "工作簿",
        "工作表",
        "单元格",
        "公式",
    ),
}
_TERM_PATTERN = re.compile(r"[A-Za-z][A-Za-z0-9_.:/-]{1,}|\d+(?:\.\d+)?|[\u4e00-\u9fff]{2,8}")
_EXPLICIT_NUMERIC_DELTA_PATTERN = re.compile(
    r"(?:\b(?:before[-\s]*(?:vs\.?|versus|and)[-\s]*after|numeric\s+(?:change|delta)|"
    r"formula\s+(?:change|delta))\b|(?:数值|公式)(?:变化|差值)|前后(?:数值|公式))",
    re.IGNORECASE,
)
_IMPLICIT_NUMERIC_DELTA_PATTERN = re.compile(
    r"(?:[%％]|\b(?:percentage\s+points?|percent|delta|difference|increase[sd]?|decrease[sd]?|"
    r"subtract(?:ed|ing)?|add(?:ed|ing)?|formula|ratio|rate)\b|(?:百分点|百分比|增(?:加|长)|"
    r"减少|差值|公式|比率|比例))",
    re.IGNORECASE,
)
_STOPWORDS = {
    "about",
    "after",
    "agent",
    "before",
    "case",
    "complete",
    "content",
    "document",
    "evidence",
    "failure",
    "file",
    "from",
    "full",
    "inspect",
    "into",
    "missing",
    "output",
    "result",
    "read",
    "show",
    "should",
    "source",
    "task",
    "that",
    "the",
    "this",
    "tool",
    "trace",
    "with",
}
_ABSENCE_CLAIM_PATTERN = re.compile(
    r"\b(?:absent|absence|missing|lacks?|without|does\s+not\s+(?:contain|include|show)|"
    r"no\s+(?:evidence|record|field|entry|occurrence))\b|(?:缺少|缺失|不存在|未包含|没有)",
    re.IGNORECASE,
)
_EXISTENCE_CLAIM_PATTERN = re.compile(
    r"\b(?:contains?|includes?|present|exists?|records?|shows?|states?)\b|(?:包含|存在|记录|显示|说明)",
    re.IGNORECASE,
)


def causal_hypothesis_semantic_id(claim: str, falsified_if: str) -> str:
    """Return a stable identity for one causal statement, independent of local labels."""
    canonical = json.dumps(
        {
            "claim": _normalized_semantics(claim),
            "falsified_if": _normalized_semantics(falsified_if),
        },
        ensure_ascii=True,
        sort_keys=True,
        separators=(",", ":"),
    )
    return f"chs:{hashlib.sha256(canonical.encode('utf-8')).hexdigest()[:24]}"


def normalize_causal_investigation(
    value: Mapping[str, Any] | None,
    *,
    failed_requirement_ids: Sequence[str] = (),
    max_requests: int = _MAX_REQUESTS,
    max_hypotheses: int = _MAX_HYPOTHESES,
    min_hypotheses: int = 1,
    min_hypotheses_per_requirement: int = 0,
    require_evidence_per_hypothesis: bool = False,
) -> dict[str, Any] | None:
    """Validate and bound a model-proposed causal investigation plan."""
    if not isinstance(value, Mapping):
        return None
    raw = value.get("causal_investigation")
    if not isinstance(raw, Mapping):
        raw = value.get("investigation")
    if not isinstance(raw, Mapping):
        raw = value
    if not isinstance(raw, Mapping):
        return None
    hypothesis_limit = max(1, min(_MAX_HYPOTHESES, int(max_hypotheses)))
    raw_hypotheses = raw.get("hypotheses")
    if not isinstance(raw_hypotheses, list) or not raw_hypotheses:
        return None

    allowed_requirements = {str(item) for item in failed_requirement_ids if str(item)}
    hypotheses: list[dict[str, Any]] = []
    hypothesis_ids: set[str] = set()
    hypothesis_semantics: set[tuple[str, str]] = set()
    raw_requests: list[tuple[str, Mapping[str, Any]]] = []
    for index, item in enumerate(raw_hypotheses[:hypothesis_limit], start=1):
        if not isinstance(item, Mapping):
            continue
        claim = str(item.get("claim", "") or "").strip()
        falsified_if = str(item.get("falsified_if", "") or "").strip()
        if not claim or not falsified_if:
            continue
        semantics = (_normalized_semantics(claim), _normalized_semantics(falsified_if))
        if semantics in hypothesis_semantics:
            continue
        hypothesis_semantics.add(semantics)
        hypothesis_id = str(item.get("hypothesis_id", "") or f"h{index}").strip()
        if not hypothesis_id or hypothesis_id in hypothesis_ids:
            hypothesis_id = f"h{index}"
        hypothesis_ids.add(hypothesis_id)
        explains = _string_list(item.get("explains_requirement_ids"))
        if allowed_requirements:
            explains = [requirement_id for requirement_id in explains if requirement_id in allowed_requirements]
        requests = item.get("evidence_requests")
        raw_hypothesis_requests = requests if isinstance(requests, list) else []
        declared_numeric_check = item.get("numeric_change_check_required")
        # A model may underestimate its own verification obligation.  Explicit
        # ``false`` therefore cannot disable a controller-detected numeric
        # before/after claim.  This is a control-plane decision, not a prompt
        # preference: textual arithmetic is never accepted as execution.
        hypothesis_text = f"{claim}\n{falsified_if}"
        numeric_change_check_required = bool(declared_numeric_check) or bool(
            _EXPLICIT_NUMERIC_DELTA_PATTERN.search(hypothesis_text)
            or _IMPLICIT_NUMERIC_DELTA_PATTERN.search(hypothesis_text)
        )
        if any(
            isinstance(request, Mapping) and request.get("operation") == "compare_numeric_change"
            for request in raw_hypothesis_requests
        ):
            numeric_change_check_required = True
        hypotheses.append(
            {
                "hypothesis_id": hypothesis_id,
                "hypothesis_semantic_id": causal_hypothesis_semantic_id(claim, falsified_if),
                "claim": claim,
                "explains_requirement_ids": explains,
                "current_support": _string_list(item.get("current_support"))[:8],
                "falsified_if": falsified_if,
                "numeric_change_check_required": numeric_change_check_required,
            }
        )
        raw_requests.extend(
            (hypothesis_id, request) for request in raw_hypothesis_requests if isinstance(request, Mapping)
        )

    top_level_requests = raw.get("evidence_requests")
    if isinstance(top_level_requests, list):
        raw_requests.extend(("", request) for request in top_level_requests if isinstance(request, Mapping))
    if len(hypotheses) < max(1, min(hypothesis_limit, int(min_hypotheses))):
        return None
    required_alternatives = max(0, min(hypothesis_limit, int(min_hypotheses_per_requirement)))
    if allowed_requirements and required_alternatives:
        hypothesis_coverage = {
            requirement_id: sum(requirement_id in hypothesis["explains_requirement_ids"] for hypothesis in hypotheses)
            for requirement_id in allowed_requirements
        }
        if any(count < required_alternatives for count in hypothesis_coverage.values()):
            return None

    request_limit = min(_MAX_REQUESTS, max(0, int(max_requests)))
    requests: list[dict[str, Any]] = []
    request_index_by_fingerprint: dict[str, int] = {}
    used_request_ids: set[str] = set()
    for default_hypothesis_id, request in raw_requests:
        normalized = _normalize_request(
            request,
            default_hypothesis_id=default_hypothesis_id,
            known_hypothesis_ids=hypothesis_ids,
            index=len(requests) + 1,
        )
        if normalized is None:
            continue
        fingerprint = _request_execution_fingerprint(normalized)
        duplicate_index = request_index_by_fingerprint.get(fingerprint)
        if duplicate_index is not None:
            existing = requests[duplicate_index]
            existing["hypothesis_ids"] = list(
                dict.fromkeys(
                    [
                        *_string_list(existing.get("hypothesis_ids")),
                        *_string_list(normalized.get("hypothesis_ids")),
                    ]
                )
            )
            continue
        if request_limit == 0 or len(requests) >= request_limit:
            # Keep scanning so later declarations can still bind another
            # hypothesis to an already retained shared probe.
            continue
        request_id = str(normalized.get("request_id", "") or f"q{len(requests) + 1}")
        if request_id in used_request_ids:
            stem = request_id
            suffix = 2
            while request_id in used_request_ids:
                request_id = f"{stem}_{suffix}"
                suffix += 1
            normalized["request_id"] = request_id
        used_request_ids.add(request_id)
        request_index_by_fingerprint[fingerprint] = len(requests)
        requests.append(normalized)

    if require_evidence_per_hypothesis:
        covered: set[str] = set()
        for request in requests:
            covered.update(_string_list(request.get("hypothesis_ids")))
        if not hypothesis_ids.issubset(covered):
            return None

    hypotheses_by_id = {item["hypothesis_id"]: item for item in hypotheses}
    for request in requests:
        if request.get("operation") != "inspect_artifact" or request.get("proof_obligation"):
            continue
        claim_text = " ".join(
            str(hypotheses_by_id[hypothesis_id].get("claim", "") or "")
            for hypothesis_id in _string_list(request.get("hypothesis_ids"))
            if hypothesis_id in hypotheses_by_id
        )
        request["proof_obligation"] = (
            "absence"
            if _ABSENCE_CLAIM_PATTERN.search(claim_text)
            else "existence"
            if claim_text and _EXISTENCE_CLAIM_PATTERN.search(claim_text)
            else "coverage"
        )

    return {
        "schema_version": 1,
        "hypotheses": hypotheses,
        "evidence_requests": requests,
        "ready_without_more_evidence": bool(raw.get("ready_without_more_evidence")) and not requests,
    }


def _request_execution_fingerprint(request: Mapping[str, Any]) -> str:
    """Identify one controller operation independently of model-local labels."""
    execution = {key: value for key, value in request.items() if key not in {"request_id", "hypothesis_ids", "purpose"}}
    return json.dumps(execution, ensure_ascii=True, sort_keys=True, separators=(",", ":"))


def execute_causal_investigation(
    case: CaseAnalysisInput,
    investigation: Mapping[str, Any],
    *,
    prior_candidate_feedback: Mapping[str, Any] | None = None,
    evidence_root: str | Path | None = None,
) -> dict[str, Any]:
    """Execute only controller-owned evidence operations for one public case."""
    trace_data = _read_json(Path(case.result_path).parent / "judge" / "normalized_trace.json")
    events = _trace_events(trace_data)
    repository_dir = _repository_dir(evidence_root)
    discovered_repository_paths: set[str] = set()
    artifact_text_cache: dict[Path, str] = {}
    results: list[dict[str, Any]] = []
    for request in investigation.get("evidence_requests", []):
        if not isinstance(request, Mapping):
            continue
        operation = str(request.get("operation", "") or "")
        if operation == "check_relation":
            evidence = _check_relation(request)
        elif operation == "compare_numeric_change":
            evidence = _compare_numeric_change(request)
        elif operation == "search_trace":
            evidence = _search_trace(events, request)
        elif operation == "read_event":
            evidence = _read_event(events, request)
        elif operation == "inspect_artifact":
            evidence = _inspect_artifact(case, request, text_cache=artifact_text_cache)
        elif operation == "read_artifact_window":
            evidence = _read_artifact_window(case, request, text_cache=artifact_text_cache)
        elif operation == "inspect_evaluation":
            evidence = _inspect_evaluation(case, request)
        elif operation == "search_repository":
            evidence = _search_repository(repository_dir, request)
            discovered_repository_paths.update(
                str(item.get("relative_path", "") or "")
                for item in evidence.get("files", [])
                if isinstance(item, Mapping) and str(item.get("relative_path", "") or "")
            )
        elif operation == "read_repository_file":
            evidence = _read_repository_file(
                repository_dir,
                request,
                discovered_paths=discovered_repository_paths,
            )
        elif operation == "compare_runs":
            evidence = _compare_runs(prior_candidate_feedback, request)
        else:
            continue
        results.append(
            {
                "request_id": str(request.get("request_id", "") or ""),
                "hypothesis_ids": _string_list(request.get("hypothesis_ids")),
                "operation": operation,
                "purpose": str(request.get("purpose", "") or ""),
                "proof_obligation": str(request.get("proof_obligation", "") or ""),
                **evidence,
            }
        )

    automatic_requests, closure_results, closure = _close_incomplete_artifact_evidence(
        case,
        investigation,
        results,
        text_cache=artifact_text_cache,
    )
    results.extend(closure_results)

    return {
        "schema_version": 1,
        "policy": {
            "controller_owned": True,
            "read_only": True,
            "arbitrary_shell_or_path_access": False,
            "display_omissions_are_not_task_agent_observations": True,
        },
        "hypotheses": list(investigation.get("hypotheses", [])),
        "request_count": len(investigation.get("evidence_requests", [])) + len(automatic_requests),
        "model_request_count": len(investigation.get("evidence_requests", [])),
        "automatic_request_count": len(automatic_requests),
        "automatic_requests": automatic_requests,
        "completed_request_count": len(results),
        "artifact_evidence_closure": closure,
        "results": results,
    }


def _close_incomplete_artifact_evidence(
    case: CaseAnalysisInput,
    investigation: Mapping[str, Any],
    initial_results: Sequence[Mapping[str, Any]],
    *,
    text_cache: dict[Path, str],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]], dict[str, Any]]:
    """Deterministically complete bounded artifact sources exposed by search.

    ``inspect_artifact`` is discovery, not proof that content is absent.  Once a
    causal request selects a physical source but returns an incomplete excerpt,
    the controller owns the mechanical continuation.  This avoids asking the
    model to rediscover an already known source/offset and gives absence claims
    continuous coverage from character zero to EOF.
    """
    existing_ids = {
        str(item.get("request_id", "") or "") for item in initial_results if str(item.get("request_id", "") or "")
    }
    # Explicit windows count toward coverage and must not be read again.
    covered_ranges: dict[str, list[tuple[int, int]]] = {}
    source_counts: dict[str, int] = {}
    for item in initial_results:
        if str(item.get("operation", "") or "") != "read_artifact_window":
            continue
        if str(item.get("availability", "") or "") != "available":
            continue
        source = str(item.get("source", "") or "")
        start = _optional_nonnegative_int(item.get("source_char_start"))
        end = _optional_nonnegative_int(item.get("source_char_end"))
        count = _optional_nonnegative_int(item.get("source_char_count"))
        if not source or start is None or end is None:
            continue
        if end >= start:
            covered_ranges.setdefault(source, []).append((start, end))
            if count is not None:
                source_counts[source] = count

    candidates: list[dict[str, Any]] = []
    candidate_by_source: dict[str, dict[str, Any]] = {}
    for item in initial_results:
        if str(item.get("operation", "") or "") != "inspect_artifact":
            continue
        if str(item.get("availability", "") or "") != "available":
            continue
        hypothesis_ids = _string_list(item.get("hypothesis_ids"))
        if not hypothesis_ids:
            continue
        if str(item.get("proof_obligation", "") or "") == "existence":
            # A physical match is a complete witness for an existence claim.
            # Reading unrelated tail content cannot strengthen that obligation.
            continue
        matches = item.get("matches", [])
        for match in matches[:1] if isinstance(matches, list) else []:
            if not isinstance(match, Mapping):
                continue
            source = str(match.get("source", "") or "")
            if not source:
                continue
            raw_spans = match.get("exact_spans", [])
            spans = raw_spans if isinstance(raw_spans, list) else []
            incomplete = any(isinstance(span, Mapping) and not bool(span.get("window_complete")) for span in spans)
            if not incomplete:
                continue
            parent_request_id = str(item.get("request_id", "") or "")
            existing_candidate = candidate_by_source.get(source)
            if existing_candidate is not None:
                existing_candidate["hypothesis_ids"] = list(
                    dict.fromkeys([*existing_candidate["hypothesis_ids"], *hypothesis_ids])
                )
                existing_candidate["parent_request_ids"] = list(
                    dict.fromkeys([*existing_candidate["parent_request_ids"], parent_request_id])
                )
                continue
            candidate = {
                "source": source,
                "logical_source": str(match.get("logical_source", "") or source),
                "hypothesis_ids": hypothesis_ids,
                "parent_request_id": parent_request_id,
                "parent_request_ids": [parent_request_id] if parent_request_id else [],
                "purpose": str(item.get("purpose", "") or ""),
            }
            candidate_by_source[source] = candidate
            candidates.append(candidate)
            if len(candidates) >= _MAX_AUTOMATIC_ARTIFACT_SOURCES:
                break
        if len(candidates) >= _MAX_AUTOMATIC_ARTIFACT_SOURCES:
            break

    automatic_requests: list[dict[str, Any]] = []
    closure_results: list[dict[str, Any]] = []
    source_records: list[dict[str, Any]] = []
    total_chars = 0
    window_budget = _MAX_AUTOMATIC_ARTIFACT_WINDOWS
    for candidate in candidates:
        source = candidate["source"]
        cursor = _continuous_prefix_end(covered_ranges.get(source, []))
        completed = bool(source_counts.get(source) is not None and cursor >= source_counts[source])
        source_windows = 0
        source_chars = 0
        last_reason = "already_covered" if completed else ""
        while not completed and window_budget > 0 and total_chars < _MAX_AUTOMATIC_ARTIFACT_CHARS:
            remaining_chars = _MAX_AUTOMATIC_ARTIFACT_CHARS - total_chars
            max_chars = min(_MAX_ARTIFACT_WINDOW_CHARS, remaining_chars)
            request_id = _unique_automatic_request_id(
                candidate["parent_request_id"],
                source,
                cursor,
                existing_ids,
            )
            request = {
                "request_id": request_id,
                "hypothesis_ids": list(candidate["hypothesis_ids"]),
                "operation": "read_artifact_window",
                "relative_path": source,
                "source_char_start": cursor,
                "max_chars": max_chars,
                "purpose": (
                    "controller-owned continuation of an incomplete artifact search; "
                    "establish continuous source coverage before reasoning about absence"
                ),
                "automatic": True,
                "parent_request_id": candidate["parent_request_id"],
                "parent_request_ids": list(candidate["parent_request_ids"]),
            }
            evidence = _read_artifact_window(case, request, text_cache=text_cache)
            result = {
                **request,
                **evidence,
            }
            automatic_requests.append(request)
            closure_results.append(result)
            existing_ids.add(request_id)
            window_budget -= 1
            source_windows += 1
            if str(evidence.get("availability", "") or "") != "available":
                last_reason = str(evidence.get("reason", "") or evidence.get("availability", "") or "unavailable")
                break
            start = _optional_nonnegative_int(evidence.get("source_char_start"))
            end = _optional_nonnegative_int(evidence.get("source_char_end"))
            count = _optional_nonnegative_int(evidence.get("source_char_count"))
            if start is None or end is None:
                last_reason = "non_contiguous_controller_window"
                break
            if count is None:
                last_reason = "non_contiguous_controller_window"
                break
            if start != cursor or end <= start:
                last_reason = "non_contiguous_controller_window"
                break
            read_chars = end - start
            total_chars += read_chars
            source_chars += read_chars
            covered_ranges.setdefault(source, []).append((start, end))
            source_counts[source] = count
            next_cursor = _continuous_prefix_end(covered_ranges[source])
            if next_cursor <= cursor:
                last_reason = "no_forward_progress"
                break
            cursor = next_cursor
            completed = cursor >= count
            last_reason = "complete" if completed else "continuation_required"
        if not completed and not last_reason:
            last_reason = "budget_exhausted"
        if not completed and (window_budget <= 0 or total_chars >= _MAX_AUTOMATIC_ARTIFACT_CHARS):
            last_reason = "budget_exhausted"
        source_records.append(
            {
                "source": source,
                "logical_source": candidate["logical_source"],
                "parent_request_id": candidate["parent_request_id"],
                "continuous_source_char_end": cursor,
                "source_char_count": source_counts.get(source),
                "window_count": source_windows,
                "read_char_count": source_chars,
                "complete": completed,
                "status": last_reason,
            }
        )

    attempted = bool(candidates)
    all_complete = attempted and all(bool(item.get("complete")) for item in source_records)
    budget_exhausted = any(item.get("status") == "budget_exhausted" for item in source_records)
    return (
        automatic_requests,
        closure_results,
        {
            "attempted": attempted,
            "status": (
                "completed"
                if all_complete
                else "budget_exhausted"
                if budget_exhausted
                else "incomplete"
                if attempted
                else "not_needed"
            ),
            "candidate_source_count": len(candidates),
            "completed_source_count": sum(bool(item.get("complete")) for item in source_records),
            "automatic_window_count": len(automatic_requests),
            "read_char_count": total_chars,
            "limits": {
                "max_sources": _MAX_AUTOMATIC_ARTIFACT_SOURCES,
                "max_windows": _MAX_AUTOMATIC_ARTIFACT_WINDOWS,
                "max_chars": _MAX_AUTOMATIC_ARTIFACT_CHARS,
                "max_chars_per_window": _MAX_ARTIFACT_WINDOW_CHARS,
            },
            "sources": source_records,
        },
    )


def _continuous_prefix_end(ranges: Sequence[tuple[int, int]]) -> int:
    cursor = 0
    for start, end in sorted(ranges):
        if start > cursor:
            break
        if end > cursor:
            cursor = end
    return cursor


def _unique_automatic_request_id(
    parent_request_id: str,
    source: str,
    start: int,
    existing_ids: set[str],
) -> str:
    stem = re.sub(r"[^a-zA-Z0-9_.-]+", "_", parent_request_id).strip("_") or "artifact"
    digest = hashlib.sha256(source.encode("utf-8")).hexdigest()[:8]
    candidate = f"{stem}.auto.{digest}.{start}"
    suffix = 2
    while candidate in existing_ids:
        candidate = f"{stem}.auto.{digest}.{start}.{suffix}"
        suffix += 1
    return candidate


def _normalize_request(
    request: Mapping[str, Any],
    *,
    default_hypothesis_id: str,
    known_hypothesis_ids: set[str],
    index: int,
) -> dict[str, Any] | None:
    operation = str(request.get("operation", "") or "").strip().casefold()
    if operation not in _ALLOWED_OPERATIONS:
        return None
    query = str(request.get("query", "") or "").strip()
    if operation in {"search_trace", "inspect_artifact", "inspect_evaluation", "search_repository"} and not query:
        return None
    requested_path = request.get("relative_path")
    if operation == "read_artifact_window" and not _safe_relative_path(requested_path):
        # Inspection results expose the selected identity as ``source``. Accept
        # that exact controller-issued identity when a refinement model feeds it
        # back, while retaining the same bounded relative-path validation.
        requested_path = request.get("source")
    if operation in {"read_artifact_window", "read_repository_file"} and not _safe_relative_path(requested_path):
        return None
    if operation == "check_relation" and not str(request.get("expression", "") or "").strip():
        return None
    if operation == "compare_numeric_change" and not all(
        str(request.get(key, "") or "").strip() for key in ("before_expression", "after_expression")
    ):
        return None
    hypothesis_ids = _string_list(request.get("hypothesis_ids"))
    if default_hypothesis_id:
        hypothesis_ids.append(default_hypothesis_id)
    hypothesis_ids = list(dict.fromkeys(item for item in hypothesis_ids if item in known_hypothesis_ids))
    normalized: dict[str, Any] = {
        "request_id": str(request.get("request_id", "") or f"q{index}").strip() or f"q{index}",
        "hypothesis_ids": hypothesis_ids,
        "operation": operation,
        "purpose": str(request.get("purpose", "") or "").strip(),
    }
    if query:
        normalized["query"] = query[:1_000]
    if operation == "inspect_artifact":
        artifact_hint = _safe_relative_path(request.get("relative_path") or request.get("source"))
        if artifact_hint:
            normalized["relative_path"] = artifact_hint
        proof_obligation = str(request.get("proof_obligation", "") or "").strip().casefold()
        if proof_obligation in {"existence", "absence", "coverage"}:
            normalized["proof_obligation"] = proof_obligation
    if operation == "check_relation":
        expected = _finite_number(request.get("expected"))
        if expected is None:
            return None
        normalized.update(
            {
                "expression": str(request.get("expression", "") or "").strip()[:200],
                "operator": str(request.get("operator", "approximately_equal") or "approximately_equal")
                .strip()
                .casefold(),
                "expected": expected,
                "tolerance": _bounded_tolerance(request.get("tolerance")),
            }
        )
    if operation == "compare_numeric_change":
        expected_delta = _finite_number(request.get("expected_delta"))
        if expected_delta is None:
            return None
        normalized.update(
            {
                "before_expression": str(request.get("before_expression", "") or "").strip()[:200],
                "after_expression": str(request.get("after_expression", "") or "").strip()[:200],
                "expected_delta": expected_delta,
                "tolerance": _bounded_tolerance(request.get("tolerance")),
            }
        )
    if operation == "search_trace":
        normalized["max_results"] = min(
            _MAX_SEARCH_RESULTS,
            max(1, _positive_int(request.get("max_results"), default=3)),
        )
    if operation == "search_repository":
        normalized["max_results"] = min(
            _MAX_SEARCH_RESULTS,
            max(1, _positive_int(request.get("max_results"), default=3)),
        )
    if operation == "read_repository_file":
        normalized["relative_path"] = _safe_relative_path(request.get("relative_path"))
    if operation == "read_artifact_window":
        normalized["relative_path"] = _safe_relative_path(requested_path)
        source_char_start = _optional_nonnegative_int(request.get("source_char_start")) or 0
        source_char_end = _optional_nonnegative_int(request.get("source_char_end"))
        requested_max_chars = request.get("max_chars")
        if requested_max_chars is None and source_char_end is not None and source_char_end > source_char_start:
            requested_max_chars = source_char_end - source_char_start
        normalized["source_char_start"] = source_char_start
        normalized["max_chars"] = min(
            _MAX_ARTIFACT_WINDOW_CHARS,
            max(1, _positive_int(requested_max_chars, default=_MAX_ARTIFACT_WINDOW_CHARS)),
        )
    if operation == "read_event":
        message_index = _optional_nonnegative_int(request.get("message_index"))
        if message_index is None:
            return None
        normalized["trace_id"] = str(request.get("trace_id", "") or "").strip()
        normalized["message_index"] = message_index
        tool_call_index = _optional_nonnegative_int(request.get("tool_call_index"))
        if tool_call_index is not None:
            normalized["tool_call_index"] = tool_call_index
    return normalized


def _check_relation(request: Mapping[str, Any]) -> dict[str, Any]:
    """Evaluate a small numeric discriminator without executing model code."""
    expression = str(request.get("expression", "") or "").strip()
    expected = _finite_number(request.get("expected"))
    operator = str(request.get("operator", "approximately_equal") or "approximately_equal").casefold()
    tolerance = _bounded_tolerance(request.get("tolerance"))
    if expected is None:
        return {"availability": "invalid", "reason": "expected_must_be_a_finite_number"}
    try:
        value = _safe_numeric_expression(expression)
    except (SyntaxError, TypeError, ValueError, ZeroDivisionError, OverflowError) as exc:
        return {
            "availability": "invalid",
            "expression": expression,
            "reason": f"invalid_numeric_expression:{type(exc).__name__}",
        }

    if operator in {"approximately_equal", "equal"}:
        holds = math.isclose(value, expected, rel_tol=0.0, abs_tol=tolerance)
    elif operator == "not_equal":
        holds = not math.isclose(value, expected, rel_tol=0.0, abs_tol=tolerance)
    elif operator == "less_than":
        holds = value < expected
    elif operator == "less_than_or_equal":
        holds = value <= expected or math.isclose(value, expected, rel_tol=0.0, abs_tol=tolerance)
    elif operator == "greater_than":
        holds = value > expected
    elif operator == "greater_than_or_equal":
        holds = value >= expected or math.isclose(value, expected, rel_tol=0.0, abs_tol=tolerance)
    else:
        return {
            "availability": "invalid",
            "expression": expression,
            "value": value,
            "reason": "unsupported_relation_operator",
        }
    return {
        "availability": "available",
        "expression": expression,
        "value": value,
        "operator": operator,
        "expected": expected,
        "tolerance": tolerance,
        "holds": holds,
    }


def _compare_numeric_change(request: Mapping[str, Any]) -> dict[str, Any]:
    """Compute after-before so formula changes are compared to their real baseline."""
    before_expression = str(request.get("before_expression", "") or "").strip()
    after_expression = str(request.get("after_expression", "") or "").strip()
    expected_delta = _finite_number(request.get("expected_delta"))
    tolerance = _bounded_tolerance(request.get("tolerance"))
    if expected_delta is None:
        return {"availability": "invalid", "reason": "expected_delta_must_be_a_finite_number"}
    try:
        before = _safe_numeric_expression(before_expression)
        after = _safe_numeric_expression(after_expression)
    except (SyntaxError, TypeError, ValueError, ZeroDivisionError, OverflowError) as exc:
        return {
            "availability": "invalid",
            "before_expression": before_expression,
            "after_expression": after_expression,
            "reason": f"invalid_numeric_expression:{type(exc).__name__}",
        }
    delta = after - before
    return {
        "availability": "available",
        "before_expression": before_expression,
        "after_expression": after_expression,
        "before_value": before,
        "after_value": after,
        "computed_delta": delta,
        "expected_delta": expected_delta,
        "tolerance": tolerance,
        "holds": math.isclose(delta, expected_delta, rel_tol=0.0, abs_tol=tolerance),
    }


def _safe_numeric_expression(expression: str) -> float:
    """Evaluate only finite numeric literals and bounded arithmetic operators."""
    if not expression or len(expression) > 200:
        raise ValueError("expression_length")
    tree = ast.parse(expression, mode="eval")
    if sum(1 for _ in ast.walk(tree)) > 40:
        raise ValueError("expression_complexity")

    def _evaluate(node: ast.AST) -> float:
        if isinstance(node, ast.Expression):
            return _evaluate(node.body)
        if isinstance(node, ast.Constant) and isinstance(node.value, (int, float)) and not isinstance(node.value, bool):
            value = float(node.value)
        elif isinstance(node, ast.UnaryOp) and isinstance(node.op, (ast.UAdd, ast.USub)):
            operand = _evaluate(node.operand)
            value = operand if isinstance(node.op, ast.UAdd) else -operand
        elif isinstance(node, ast.BinOp) and isinstance(node.op, (ast.Add, ast.Sub, ast.Mult, ast.Div)):
            left = _evaluate(node.left)
            right = _evaluate(node.right)
            if isinstance(node.op, ast.Add):
                value = left + right
            elif isinstance(node.op, ast.Sub):
                value = left - right
            elif isinstance(node.op, ast.Mult):
                value = left * right
            else:
                value = left / right
        else:
            raise ValueError("unsupported_expression_node")
        if not math.isfinite(value) or abs(value) > 1e100:
            raise ValueError("non_finite_or_unbounded_result")
        return value

    return _evaluate(tree)


def _finite_number(value: Any) -> float | None:
    if isinstance(value, bool) or not isinstance(value, (int, float)):
        return None
    number = float(value)
    return number if math.isfinite(number) else None


def _bounded_tolerance(value: Any) -> float:
    number = _finite_number(value)
    if number is None:
        return 1e-9
    return min(max(number, 0.0), 1.0)


def _trace_events(trace_data: Mapping[str, Any]) -> list[dict[str, Any]]:
    events: list[dict[str, Any]] = []
    traces = trace_data.get("traces")
    if not isinstance(traces, list):
        return events
    for trace in traces:
        if not isinstance(trace, Mapping):
            continue
        trace_id = str(trace.get("trace_id", "") or "")
        messages = trace.get("messages")
        if not isinstance(messages, list):
            continue
        for sequence, message in enumerate(messages):
            if not isinstance(message, Mapping):
                continue
            raw_index = message.get("message_index", sequence)
            message_index = _optional_nonnegative_int(raw_index)
            if message_index is None:
                message_index = sequence
            tool_calls = [dict(call) for call in message.get("tool_calls", []) if isinstance(call, Mapping)]
            events.append(
                {
                    "trace_id": trace_id,
                    "message_index": message_index,
                    "step_pointer": str(message.get("step_pointer", "") or ""),
                    "role": str(message.get("role", "") or ""),
                    "content": str(message.get("content", "") or ""),
                    "tool_calls": tool_calls,
                }
            )
    return events


def _search_trace(events: list[dict[str, Any]], request: Mapping[str, Any]) -> dict[str, Any]:
    query = str(request.get("query", "") or "")
    terms = _query_terms(query)
    ranked: list[tuple[int, int, dict[str, Any]]] = []
    for sequence, event in enumerate(events):
        searchable = json.dumps(event, ensure_ascii=False, separators=(",", ":")).casefold()
        matched = [term for term in terms if term in searchable]
        if matched:
            ranked.append((len(set(matched)), sequence, event))
    limit = min(_MAX_SEARCH_RESULTS, max(1, _positive_int(request.get("max_results"), default=3)))
    selected = sorted(ranked, key=lambda item: (-item[0], item[1]))[:limit]
    return {
        "availability": "available" if selected else "not_found",
        "query": query,
        "matched_event_count": len(ranked),
        "events": [_search_event_view(event, terms) for _, _, event in selected],
    }


def _read_event(events: list[dict[str, Any]], request: Mapping[str, Any]) -> dict[str, Any]:
    trace_id = str(request.get("trace_id", "") or "")
    message_index = _optional_nonnegative_int(request.get("message_index"))
    matches = [
        event
        for event in events
        if event["message_index"] == message_index and (not trace_id or event["trace_id"] == trace_id)
    ]
    if not matches:
        return {
            "availability": "not_found",
            "trace_id": trace_id,
            "message_index": message_index,
        }
    if not trace_id and len(matches) > 1:
        return {
            "availability": "ambiguous",
            "reason": "trace_id_required_for_non_unique_message_index",
            "message_index": message_index,
            "candidate_trace_ids": sorted({str(item.get("trace_id", "") or "") for item in matches}),
        }
    event = matches[0]
    tool_call_index = _optional_nonnegative_int(request.get("tool_call_index"))
    calls = event["tool_calls"]
    if tool_call_index is not None:
        calls = [calls[tool_call_index]] if tool_call_index < len(calls) else []
    return {
        "availability": "available",
        "event": {
            "trace_id": event["trace_id"],
            "message_index": event["message_index"],
            "step_pointer": event["step_pointer"],
            "role": event["role"],
            "content": _bounded_exact_text(event["content"], _MAX_EVENT_CHARS),
            "tool_calls": [
                {
                    "tool_call_index": index,
                    "name": str(call.get("name", "") or ""),
                    "input": _bounded_exact_text(str(call.get("input", "") or ""), _MAX_EVENT_CHARS),
                    "output": _bounded_exact_text(str(call.get("output", "") or ""), _MAX_EVENT_CHARS),
                    "error": _bounded_exact_text(str(call.get("error", "") or ""), 4_000),
                }
                for index, call in enumerate(calls)
            ],
        },
    }


def _inspect_artifact(
    case: CaseAnalysisInput,
    request: Mapping[str, Any],
    *,
    text_cache: dict[Path, str] | None = None,
) -> dict[str, Any]:
    """Search only physically materialized task artifacts.

    Evaluation metadata is deliberately excluded.  Previously a criterion or
    judge explanation could satisfy an ``inspect_artifact`` request even when
    the source document/spreadsheet was absent, which turned a repeated outcome
    description into apparent causal evidence.
    """
    query = str(request.get("query", "") or "")
    artifact_hint = str(request.get("relative_path", "") or "").strip()
    purpose = str(request.get("purpose", "") or "").strip()
    selection_query = f"{query} {artifact_hint} {purpose}".strip()
    terms = _query_terms(query)
    identity_terms = _query_terms(f"{artifact_hint} {purpose}".strip())
    case_dir = _windows_long_path(Path(case.result_path).parent)
    sources: list[tuple[str, str, str]] = []
    artifacts_dir = case_dir / "artifacts"
    if artifacts_dir.is_dir():
        aliases = _artifact_path_aliases(case, artifacts_dir)
        files = [
            path
            for path in sorted(artifacts_dir.rglob("*"))
            if path.is_file() and path.suffix.casefold() in _TEXT_ARTIFACT_SUFFIXES
        ][:_MAX_ARTIFACT_FILES]
        for path in files:
            resolved = path.resolve()
            if not resolved.is_relative_to(case_dir):
                continue
            source = f"artifacts/{path.relative_to(artifacts_dir).as_posix()}"
            sources.append((source, aliases.get(resolved, source), _read_text(path, _MAX_ARTIFACT_FILE_CHARS)))
        structured_files = [
            path
            for path in sorted(artifacts_dir.rglob("*"))
            if path.is_file() and path.suffix.casefold() in _STRUCTURED_ARTIFACT_SUFFIXES
        ]
        structured_files = _select_structured_artifact_files(structured_files, selection_query, aliases=aliases)
        for path in structured_files:
            resolved = path.resolve()
            if not resolved.is_relative_to(case_dir):
                continue
            if text_cache is not None and resolved in text_cache:
                text = text_cache[resolved]
            else:
                text = _structured_artifact_text(path)
                if text_cache is not None:
                    text_cache[resolved] = text
            if text:
                source = f"artifacts/{path.relative_to(artifacts_dir).as_posix()}"
                sources.append((source, aliases.get(resolved, source), text))

    if not sources:
        return {
            "availability": "not_available",
            "reason": "physical_artifact_snapshot_not_available",
            "query": query,
            "matches": [],
        }

    if artifact_hint:
        hint_terms = _query_terms(artifact_hint)
        identity_scores = [
            (_query_match_score(logical_source, hint_terms), logical_source) for _, logical_source, _ in sources
        ]
        best_identity_score = max((score for score, _ in identity_scores), default=0.0)
        if best_identity_score > 0:
            best_logical_sources = {
                logical_source for score, logical_source in identity_scores if score == best_identity_score
            }
            sources = [item for item in sources if item[1] in best_logical_sources]

    matches: list[tuple[float, str, str, str]] = []
    for source, logical_source, text in sources:
        lowered = text.casefold()
        matched = [term for term in terms if term in lowered]
        path_matched = [term for term in terms if term in logical_source.casefold()]
        if matched or path_matched:
            score = (
                _query_match_score(text, terms)
                + 2.0 * _query_match_score(logical_source, terms)
                + 4.0 * _query_match_score(logical_source, identity_terms)
            )
            matches.append((score, source, logical_source, text))
    selected = sorted(matches, key=lambda item: (-item[0], item[1]))[:_MAX_SEARCH_RESULTS]
    return {
        "availability": "available" if selected else "not_found",
        "query": query,
        "matches": [
            {
                "source": source,
                "logical_source": logical_source,
                "exact_spans": _exact_match_spans(text, terms, max_spans=3) or _leading_text_span(text),
            }
            for _, source, logical_source, text in selected
        ],
    }


def _artifact_path_aliases(case: CaseAnalysisInput, artifacts_dir: Path) -> dict[Path, str]:
    """Map bounded snapshot paths back to their original logical names."""
    snapshot = case.evaluation_metadata.get("analysis_artifact_snapshot", {})
    rows = snapshot.get("files", []) if isinstance(snapshot, Mapping) else []
    aliases: dict[Path, str] = {}
    for row in rows if isinstance(rows, list) else []:
        if not isinstance(row, Mapping):
            continue
        stored = _safe_relative_path(row.get("path"))
        logical = str(row.get("source_path", "") or "").strip()
        if not stored or not logical:
            continue
        for candidate in (artifacts_dir / stored, artifacts_dir / "workspace" / stored):
            resolved = candidate.resolve()
            if resolved.is_relative_to(artifacts_dir) and resolved.is_file():
                aliases[resolved] = logical
                break
    return aliases


def _select_structured_artifact_files(
    files: Sequence[Path],
    query: str,
    *,
    aliases: Mapping[Path, str] | None = None,
) -> list[Path]:
    """Bound structured parsing and prefer file types named by the request."""
    lowered_query = query.casefold()
    hinted_suffixes = {
        suffix
        for suffix, hints in _STRUCTURED_QUERY_SUFFIX_HINTS.items()
        if any(hint in lowered_query for hint in hints)
    }
    candidates = [path for path in files if not hinted_suffixes or path.suffix.casefold() in hinted_suffixes]
    if len(candidates) <= _MAX_STRUCTURED_ARTIFACTS_PER_REQUEST:
        return candidates

    query_terms = _query_terms(query)
    ranked = sorted(
        candidates,
        key=lambda path: (
            -_query_match_score(
                f"{path.as_posix()} {(aliases or {}).get(path.resolve(), '')}",
                query_terms,
            ),
            path.stat().st_size,
            path.as_posix(),
        ),
    )
    return ranked[:_MAX_STRUCTURED_ARTIFACTS_PER_REQUEST]


def _read_artifact_window(
    case: CaseAnalysisInput,
    request: Mapping[str, Any],
    *,
    text_cache: dict[Path, str] | None = None,
) -> dict[str, Any]:
    """Read one exact, bounded window from a previously named task artifact."""
    relative_path = _safe_relative_path(request.get("relative_path"))
    if relative_path.startswith("artifacts/"):
        relative_path = relative_path.removeprefix("artifacts/")
    if not relative_path:
        return {"availability": "invalid", "reason": "invalid_relative_path"}
    case_dir = _windows_long_path(Path(case.result_path).parent)
    artifacts_dir = (case_dir / "artifacts").resolve()
    aliases = _artifact_path_aliases(case, artifacts_dir)
    path = (artifacts_dir / relative_path).resolve()
    logical_source = aliases.get(path)
    if not path.is_relative_to(artifacts_dir) or not path.is_file():
        workspace_path = (artifacts_dir / "workspace" / relative_path).resolve()
        if workspace_path.is_relative_to(artifacts_dir) and workspace_path.is_file():
            path = workspace_path
            logical_source = aliases.get(path)
        else:
            requested = relative_path.removeprefix("workspace/")
            logical_matches = [
                (stored, logical)
                for stored, logical in aliases.items()
                if logical.replace("\\", "/").removeprefix("workspace/") == requested
            ]
            if not logical_matches:
                logical_matches = _unambiguous_logical_artifact_matches(requested, aliases)
            if len(logical_matches) != 1:
                return {"availability": "not_found", "relative_path": relative_path}
            path, logical_source = logical_matches[0]
    if not path.is_relative_to(artifacts_dir) or not path.is_file():
        return {"availability": "not_found", "relative_path": relative_path}
    suffix = path.suffix.casefold()
    if suffix in _STRUCTURED_ARTIFACT_SUFFIXES:
        if text_cache is not None and path in text_cache:
            content = text_cache[path]
        else:
            content = _structured_artifact_text(path)
            if text_cache is not None:
                text_cache[path] = content
    elif suffix in _TEXT_ARTIFACT_SUFFIXES:
        content = _read_text(path, _MAX_ARTIFACT_FILE_CHARS)
    else:
        return {"availability": "unsupported", "relative_path": relative_path}
    if not content:
        return {"availability": "not_available", "relative_path": relative_path}
    start = min(
        len(content),
        _optional_nonnegative_int(request.get("source_char_start")) or 0,
    )
    max_chars = min(
        _MAX_ARTIFACT_WINDOW_CHARS,
        max(1, _positive_int(request.get("max_chars"), default=_MAX_ARTIFACT_WINDOW_CHARS)),
    )
    end = min(len(content), start + max_chars)
    physical_relative_path = path.relative_to(artifacts_dir).as_posix()
    return {
        "availability": "available",
        "source": f"artifacts/{physical_relative_path}",
        "logical_source": logical_source or aliases.get(path, f"artifacts/{physical_relative_path}"),
        "source_char_start": start,
        "source_char_end": end,
        "source_char_count": len(content),
        "text": content[start:end],
        "window_complete": start == 0 and end == len(content),
        "next_source_char_start": end if end < len(content) else None,
        "omission_origin": "controller_read_window" if start or end < len(content) else "none",
    }


def _unambiguous_logical_artifact_matches(
    requested: str,
    aliases: Mapping[Path, str],
) -> list[tuple[Path, str]]:
    """Recover a named artifact when presentation-only path encoding changed."""
    requested_terms = {term for term in _query_terms(Path(requested).name) if len(term) >= 3}
    if len(requested_terms) < 2:
        return []
    ranked: list[tuple[int, Path, str]] = []
    for stored, logical in aliases.items():
        logical_terms = {term for term in _query_terms(Path(logical).name) if len(term) >= 3}
        overlap = len(requested_terms & logical_terms)
        if overlap:
            ranked.append((overlap, stored, logical))
    if not ranked:
        return []
    best = max(score for score, _, _ in ranked)
    minimum = max(2, (len(requested_terms) + 1) // 2)
    winners = [(stored, logical) for score, stored, logical in ranked if score == best and score >= minimum]
    return winners if len(winners) == 1 else []


def _inspect_evaluation(case: CaseAnalysisInput, request: Mapping[str, Any]) -> dict[str, Any]:
    """Search evaluator-owned result metadata without calling it an artifact."""
    query = str(request.get("query", "") or "")
    terms = _query_terms(query)
    sources = [
        ("case.evaluation_metadata", json.dumps(case.evaluation_metadata, ensure_ascii=False, indent=2)),
        ("case.result", _read_text(Path(case.result_path), _MAX_ARTIFACT_FILE_CHARS)),
    ]
    matches: list[tuple[int, str, str]] = []
    for source, content in sources:
        lowered = content.casefold()
        matched = [term for term in terms if term in lowered]
        if matched:
            matches.append((len(set(matched)), source, content))
    selected = sorted(matches, key=lambda item: (-item[0], item[1]))[:_MAX_SEARCH_RESULTS]
    return {
        "availability": "available" if selected else "not_found",
        "evidence_class": "evaluation_metadata",
        "query": query,
        "matches": [
            {
                "source": source,
                "exact_spans": _exact_match_spans(content, terms, max_spans=3),
            }
            for _, source, content in selected
        ],
    }


def _repository_dir(evidence_root: str | Path | None) -> Path | None:
    if evidence_root is None:
        return None
    root = Path(evidence_root).expanduser().resolve()
    repository = (root / "repository").resolve()
    return repository if repository.is_dir() and repository.is_relative_to(root) else None


def _search_repository(repository_dir: Path | None, request: Mapping[str, Any]) -> dict[str, Any]:
    if repository_dir is None:
        return {"availability": "not_available", "reason": "repository_snapshot_not_available"}
    query = str(request.get("query", "") or "")
    terms = _query_terms(query)
    ranked: list[tuple[int, str, str]] = []
    for path in _repository_text_files(repository_dir):
        text = _read_text(path, _MAX_ARTIFACT_FILE_CHARS)
        lowered = text.casefold()
        matched = [term for term in terms if term in lowered]
        if matched:
            relative = path.relative_to(repository_dir).as_posix()
            ranked.append((len(set(matched)), relative, text))
    limit = min(_MAX_SEARCH_RESULTS, max(1, _positive_int(request.get("max_results"), default=3)))
    selected = sorted(ranked, key=lambda item: (-item[0], item[1]))[:limit]
    return {
        "availability": "available" if selected else "not_found",
        "query": query,
        "matched_file_count": len(ranked),
        "files": [
            {"relative_path": relative, "exact_spans": _exact_match_spans(text, terms, max_spans=3)}
            for _, relative, text in selected
        ],
    }


def _read_repository_file(
    repository_dir: Path | None,
    request: Mapping[str, Any],
    *,
    discovered_paths: set[str],
) -> dict[str, Any]:
    if repository_dir is None:
        return {"availability": "not_available", "reason": "repository_snapshot_not_available"}
    relative_path = _safe_relative_path(request.get("relative_path"))
    if not relative_path:
        return {"availability": "invalid", "reason": "invalid_relative_path"}
    if relative_path not in discovered_paths:
        return {
            "availability": "invalid",
            "reason": "relative_path_not_returned_by_prior_search_repository",
            "relative_path": relative_path,
        }
    path = (repository_dir / relative_path).resolve()
    if not path.is_relative_to(repository_dir) or not path.is_file():
        return {"availability": "not_found", "relative_path": relative_path}
    if path.suffix.casefold() not in _TEXT_ARTIFACT_SUFFIXES | {
        ".c",
        ".cc",
        ".cpp",
        ".go",
        ".h",
        ".hpp",
        ".java",
        ".js",
        ".jsx",
        ".py",
        ".rs",
        ".sh",
        ".ts",
        ".tsx",
    }:
        return {"availability": "unsupported", "relative_path": relative_path}
    return {
        "availability": "available",
        "relative_path": relative_path,
        "content": _bounded_exact_text(_read_text(path, _MAX_ARTIFACT_FILE_CHARS), _MAX_EVENT_CHARS),
    }


def _repository_text_files(repository_dir: Path) -> list[Path]:
    suffixes = _TEXT_ARTIFACT_SUFFIXES | {
        ".c",
        ".cc",
        ".cpp",
        ".go",
        ".h",
        ".hpp",
        ".java",
        ".js",
        ".jsx",
        ".py",
        ".rs",
        ".sh",
        ".ts",
        ".tsx",
    }
    return [
        path
        for path in sorted(repository_dir.rglob("*"))
        if path.is_file() and path.suffix.casefold() in suffixes and path.resolve().is_relative_to(repository_dir)
    ][:_MAX_REPOSITORY_FILES]


def _safe_relative_path(value: Any) -> str:
    text = str(value or "").strip().replace("\\", "/")
    if not text or text.startswith("/") or re.match(r"^[A-Za-z]:", text):
        return ""
    parts = [part for part in text.split("/") if part not in {"", "."}]
    if not parts or any(part == ".." for part in parts):
        return ""
    return "/".join(parts)


def _normalized_semantics(value: str) -> str:
    return " ".join(re.findall(r"[a-z0-9_]+|[\u4e00-\u9fff]", value.casefold()))


def _structured_artifact_text(path: Path) -> str:
    try:
        suffix = path.suffix.casefold()
        if suffix == ".xlsx":
            return _xlsx_text(path)
        if suffix == ".docx":
            return _docx_text(path)
        if suffix == ".pdf":
            return _pdf_text(path)
        if suffix == ".pptx":
            return _pptx_text(path)
    except Exception:  # malformed optional artifacts are unavailable evidence
        return ""
    return ""


def _xlsx_text(path: Path) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(path, read_only=True, data_only=False)
    lines: list[str] = []
    total_chars = 0
    cell_count = 0
    try:
        for worksheet in workbook.worksheets:
            lines.append(f"[sheet:{worksheet.title}]")
            for row in worksheet.iter_rows(
                max_row=min(worksheet.max_row, 2_000),
                max_col=min(worksheet.max_column, 200),
            ):
                for cell in row:
                    if cell.value is not None:
                        line = f"{cell.coordinate}={cell.value}"
                        lines.append(line)
                        total_chars += len(line)
                        cell_count += 1
                        if total_chars >= _MAX_ARTIFACT_FILE_CHARS or cell_count >= _MAX_STRUCTURED_CELLS:
                            return "\n".join(lines)
    finally:
        workbook.close()
    return "\n".join(lines)


def _docx_text(path: Path) -> str:
    from docx import Document

    document = Document(path)
    lines: list[str] = []
    total_chars = 0
    for paragraph in document.paragraphs:
        if not paragraph.text:
            continue
        lines.append(paragraph.text)
        total_chars += len(paragraph.text)
        if total_chars >= _MAX_ARTIFACT_FILE_CHARS:
            return "\n".join(lines)[:_MAX_ARTIFACT_FILE_CHARS]
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            values = [cell.text for cell in row.cells]
            line = f"[table:{table_index}:row:{row_index}] " + " | ".join(values)
            lines.append(line)
            total_chars += len(line)
            if total_chars >= _MAX_ARTIFACT_FILE_CHARS:
                return "\n".join(lines)[:_MAX_ARTIFACT_FILE_CHARS]
    return "\n".join(lines)[:_MAX_ARTIFACT_FILE_CHARS]


def _pdf_text(path: Path) -> str:
    import pdfplumber

    lines: list[str] = []
    with pdfplumber.open(path) as document:
        for page_index, page in enumerate(document.pages[:_MAX_STRUCTURED_PAGES], start=1):
            lines.append(f"[page:{page_index}]")
            lines.append(page.extract_text() or "")
            if sum(len(item) for item in lines) >= _MAX_ARTIFACT_FILE_CHARS:
                break
    return "\n".join(lines)[:_MAX_ARTIFACT_FILE_CHARS]


def _pptx_text(path: Path) -> str:
    lines: list[str] = []
    with zipfile.ZipFile(path) as archive:
        slide_names = sorted(name for name in archive.namelist() if re.fullmatch(r"ppt/slides/slide\d+\.xml", name))[
            :_MAX_STRUCTURED_PAGES
        ]
        for slide_name in slide_names:
            root = ElementTree.fromstring(archive.read(slide_name))
            texts = [element.text or "" for element in root.iter() if element.tag.endswith("}t")]
            lines.append(f"[{slide_name}] " + " ".join(texts))
    return "\n".join(lines)[:_MAX_ARTIFACT_FILE_CHARS]


def _compare_runs(
    prior_candidate_feedback: Mapping[str, Any] | None,
    request: Mapping[str, Any],
) -> dict[str, Any]:
    if not prior_candidate_feedback:
        return {"availability": "not_available", "reason": "no_paired_candidate_feedback"}
    rendered = json.dumps(prior_candidate_feedback, ensure_ascii=False, indent=2)
    query = str(request.get("query", "") or "")
    terms = _query_terms(query) if query else []
    return {
        "availability": "available",
        "query": query,
        "paired_feedback": (
            {"exact_spans": _exact_match_spans(rendered, terms, max_spans=5)}
            if terms
            else _bounded_exact_text(rendered, _MAX_EVENT_CHARS)
        ),
    }


def _search_event_view(event: Mapping[str, Any], terms: Sequence[str]) -> dict[str, Any]:
    return {
        "trace_id": event.get("trace_id"),
        "message_index": event.get("message_index"),
        "step_pointer": event.get("step_pointer"),
        "role": event.get("role"),
        "content_spans": _exact_match_spans(str(event.get("content", "") or ""), terms, max_spans=2),
        "tool_calls": [
            {
                "tool_call_index": index,
                "name": str(call.get("name", "") or ""),
                "input_spans": _exact_match_spans(str(call.get("input", "") or ""), terms, max_spans=2),
                "output_spans": _exact_match_spans(str(call.get("output", "") or ""), terms, max_spans=3),
                "error_spans": _exact_match_spans(str(call.get("error", "") or ""), terms, max_spans=2),
            }
            for index, call in enumerate(event.get("tool_calls", []))
        ],
    }


def _exact_match_spans(
    text: str,
    terms: Sequence[str],
    *,
    max_spans: int,
    span_chars: int = 2_000,
) -> list[dict[str, Any]]:
    if not text or not terms:
        return []
    lowered = text.casefold()
    candidates: list[tuple[float, int, int, list[str]]] = []
    for term in terms:
        start_at = 0
        for _ in range(16):
            position = lowered.find(term, start_at)
            if position < 0:
                break
            start = max(0, position - span_chars // 3)
            end = min(len(text), start + span_chars)
            start = max(0, end - span_chars)
            window = lowered[start:end]
            matched_terms = [candidate for candidate in terms if candidate in window]
            score = sum(_term_weight(candidate) for candidate in set(matched_terms))
            score += min(1.0, len(matched_terms) / 10)
            candidates.append((score, start, end, matched_terms))
            start_at = position + max(1, len(term))
    spans: list[dict[str, Any]] = []
    occupied: list[tuple[int, int]] = []
    for _, start, end, matched_terms in sorted(candidates, key=lambda item: (-item[0], item[1])):
        if any(start < old_end and end > old_start for old_start, old_end in occupied):
            continue
        occupied.append((start, end))
        spans.append(
            {
                "source_char_start": start,
                "source_char_end": end,
                "source_char_count": len(text),
                "matched_term": matched_terms[0],
                "matched_terms": list(dict.fromkeys(matched_terms)),
                "text": text[start:end],
                "window_complete": start == 0 and end == len(text),
                "omission_origin": "controller_search_window" if start or end < len(text) else "none",
            }
        )
        if len(spans) >= max_spans:
            break
    return spans


def _leading_text_span(text: str, *, span_chars: int = 2_000) -> list[dict[str, Any]]:
    if not text:
        return []
    end = min(len(text), span_chars)
    return [
        {
            "source_char_start": 0,
            "source_char_end": end,
            "source_char_count": len(text),
            "matched_term": "logical_source",
            "matched_terms": ["logical_source"],
            "text": text[:end],
            "window_complete": end == len(text),
            "omission_origin": "controller_search_window" if end < len(text) else "none",
        }
    ]


def _query_match_score(text: str, terms: Sequence[str]) -> float:
    lowered = text.casefold()
    return sum(_term_weight(term) for term in set(terms) if term in lowered)


def _term_weight(term: str) -> float:
    weight = 1.0 + min(len(term), 24) / 24
    if any(character.isdigit() for character in term):
        weight += 1.5
    if any(character in term for character in "_./:-"):
        weight += 0.5
    return weight


def _bounded_exact_text(text: str, limit: int) -> dict[str, Any]:
    if len(text) <= limit:
        return {
            "text": text,
            "source_char_count": len(text),
            "complete": True,
            "omission_origin": "none",
        }
    return {
        "text": text[:limit],
        "source_char_count": len(text),
        "complete": False,
        "omission_origin": "controller_bound",
        "omitted_source_chars": len(text) - limit,
    }


def _query_terms(query: str) -> list[str]:
    terms: list[str] = []
    seen: set[str] = set()
    for match in _TERM_PATTERN.findall(query):
        normalized = match.casefold()
        if normalized in _STOPWORDS or normalized in seen:
            continue
        seen.add(normalized)
        terms.append(normalized)
        if len(terms) >= 24:
            break
    return terms


def _read_json(path: Path) -> dict[str, Any]:
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {}
    return value if isinstance(value, dict) else {}


def _read_text(path: Path, limit: int) -> str:
    try:
        with path.open("r", encoding="utf-8", errors="replace") as stream:
            return stream.read(limit)
    except OSError:
        return ""


def _windows_long_path(path: Path) -> Path:
    """Return an extended Windows path for deep evaluation artifact trees."""
    resolved = str(path.resolve(strict=False))
    if os.name != "nt" or resolved.startswith("\\\\?\\"):
        return Path(resolved)
    if resolved.startswith("\\\\"):
        return Path("".join(("\\\\?\\UNC\\", resolved.lstrip("\\"))))
    return Path("".join(("\\\\?\\", resolved)))


def _string_list(value: Any) -> list[str]:
    if not isinstance(value, (list, tuple, set)):
        return []
    return [str(item).strip() for item in value if str(item).strip()]


def _positive_int(value: Any, *, default: int) -> int:
    try:
        parsed = int(value)
    except (TypeError, ValueError):
        return default
    return parsed if parsed > 0 else default


def _optional_nonnegative_int(value: Any) -> int | None:
    try:
        parsed = int(value)
    except (TypeError, ValueError):
        return None
    return parsed if parsed >= 0 else None


__all__ = [
    "execute_causal_investigation",
    "normalize_causal_investigation",
]
